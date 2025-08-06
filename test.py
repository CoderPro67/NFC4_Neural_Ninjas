import os
import json
import time
import io
import docx
from dotenv import load_dotenv
from tqdm import tqdm
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

# Load environment variables
load_dotenv()

# Define compliance modes
COMPLIANCE_MAP = {
    "gdpr": ["names", "emails", "phones", "physical mailing addresses", "IP addresses", "social security numbers", "passport numbers"],
    "hipaa": ["names", "dates", "phone numbers", "emails", "social security numbers", "medical record numbers"],
    "dpdp": ["names", "emails", "phones", "addresses", "Aadhaar numbers", "PAN numbers", "financial data"]
}

# Initialize Gemini model
try:
    llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0)
except Exception as e:
    print(f"❌ Error initializing Gemini model: {e}")
    llm = None

# Define prompt template
prompt_template = ChatPromptTemplate.from_messages([
    ("system", """
You are an AI assistant trained to identify and redact Personally Identifiable Information (PII) from documents.

Instructions:
- Detect PII types: {entity_types}
- Replace each PII instance with "[REDACTED]"
- Preserve the original structure, formatting, and line breaks
- Do not remove or alter non-PII content
- Return only the redacted text. No explanation or summary.
- also follow the compliance map to redact PAN numbers,aaddhar and bank account numbers
Now process the following document:
"""),
    ("human", "Document Text:\n---\n{document_text}\n---\n\nRedacted Output:")
])


redaction_chain = prompt_template | llm | StrOutputParser() if llm else None

def load_document(file_path):
    """Load document content based on file type."""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == '.pdf':
        try:
            reader = PdfReader(file_path)
            text = ""
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {i + 1} ---\n{page_text}\n"
            return text.strip()
        except Exception as e:
            raise RuntimeError(f"Error reading PDF: {e}")

    elif ext == '.docx':
        try:
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            raise RuntimeError(f"Error reading DOCX: {e}")

    elif ext == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise RuntimeError(f"Error reading TXT: {e}")

    elif ext == '.json':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.dumps(json.load(f), indent=2)
        except Exception as e:
            raise RuntimeError(f"Error reading JSON: {e}")

    else:
        raise ValueError(f"Unsupported file type: {ext}")



def apply_pdf_redaction(input_path, output_path, redaction_boxes):
    """Overlay black rectangles on detected PII in PDF."""
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()

        for page_num, page in enumerate(reader.pages):
            writer.add_page(page)

            # Create overlay canvas
            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=letter)
            can.setFillColorRGB(0, 0, 0)

            # Draw redaction boxes for this page
            for box in redaction_boxes:
                if box.get('page') == page_num:
                    x, y = box['x'], box['y']
                    w, h = box['width'], box['height']
                    can.rect(x, y, w, h, fill=1)

            can.save()
            packet.seek(0)
            overlay = PdfReader(packet)
            writer.pages[page_num].merge_page(overlay.pages[0])

        # Save final redacted PDF
        with open(output_path, 'wb') as f:
            writer.write(f)

        print(f"✅ Redacted PDF saved to: {output_path}")

    except Exception as e:
        print(f"❌ Error applying PDF redaction: {e}")


def save_document(output_path, redacted_text, input_format, original_path):
    """Save redacted output to file."""
    print(f"💾 Saving redacted file to: {output_path}")
    try:
        if not redacted_text.strip():
            raise ValueError("❌ Redacted text is empty. Nothing to save.")

        if input_format == '.pdf':
            # Save redacted text as a simple PDF
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            for line in redacted_text.split('\n'):
                pdf.multi_cell(0, 10, line)
            pdf.output(output_path)

        elif input_format == '.docx':
            import docx
            doc = docx.Document()
            for line in redacted_text.split('\n'):
                doc.add_paragraph(line)
            doc.save(output_path)

        elif input_format in ['.txt', '.json']:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(redacted_text)

        else:
            print(f"⚠️ Unsupported format '{input_format}'. Saving as .txt.")
            with open(output_path + '.txt', 'w', encoding='utf-8') as f:
                f.write(redacted_text)

    except Exception as e:
        print(f"❌ Error saving output file: {e}")




def main():
    if not redaction_chain:
        print("❌ Redaction chain not initialized.")
        return

    input_file = input("📄 Enter the path to the input file (pdf, docx, txt, json): ").strip()
    if not os.path.exists(input_file):
        print(f"❌ File '{input_file}' not found.")
        return

    _, ext = os.path.splitext(input_file)
    ext = ext.lower()
    if ext not in COMPLIANCE_MAP.keys() and ext not in ['.pdf', '.docx', '.txt', '.json']:
        print(f"❌ Unsupported file type '{ext}'.")
        return

    print(f"ℹ️ Available compliance modes: {', '.join(COMPLIANCE_MAP.keys())}")
    compliance_mode = input("🔒 Enter compliance mode (gdpr, hipaa, dpdp): ").strip().lower()
    if compliance_mode not in COMPLIANCE_MAP:
        print(f"❌ Invalid mode '{compliance_mode}'.")
        return

    entity_types = COMPLIANCE_MAP[compliance_mode]
    print(f"🔍 Redacting: {', '.join(entity_types)}")

    try:
        print(f"\n📄 Loading document: {input_file}")
        document_text = load_document(input_file)
    except Exception as e:
        print(f"❌ Error loading document: {e}")
        return

    print("🤖 Sending to Gemini for redaction...")
    with tqdm(total=100, desc="   Processing") as pbar:
        pbar.update(20)
        time.sleep(0.5)
        try:
            redacted_text = redaction_chain.invoke({
                "document_text": document_text,
                "entity_types": ", ".join(entity_types)
            })
            pbar.update(80)
        except Exception as e:
            print(f"\n❌ API error: {e}")
            return

    print("✅ Redaction complete.")
    base, _ = os.path.splitext(input_file)
    output_path = f"{base}_redacted{ext}"

    try:
        save_document(output_path, redacted_text.strip(), ext, input_file)
        print(f"\n🎉 Success! File saved to: {output_path}")
    except Exception as e:
        print(f"❌ Error saving file: {e}")

if __name__ == "__main__":
    main()
